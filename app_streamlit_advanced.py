import os
from dotenv import load_dotenv
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from fp_a_agent import run_fp_a_agent
import io
from datetime import datetime
from data_analyzer import DataAnalyzer
import json
import re

# Load .env (GOOGLE_API_KEY)
load_dotenv()

st.set_page_config(page_title="FP&A AI Agent", layout="wide")
st.title("📊 FP&A AI Agent – CFO-Grade FP&A with Data Analytics")
st.markdown("""
    <div style='text-align: center; padding: 1rem 0 2rem 0;'>
        <p style='font-size: 1.2rem; color: #B0BEC5; font-weight: 400;'>
            🚀 Powered by Google Gemini AI | 📊 Advanced Analytics | 📈 Interactive Visualizations
        </p>
    </div>
""", unsafe_allow_html=True)

# Custom CSS for professional styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    .main { font-family: 'Inter', sans-serif; }
    h1 { font-weight: 700 !important; font-size: 2.5rem !important; background: linear-gradient(120deg, #1E88E5 0%, #00BCD4 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; margin-bottom: 1rem !important; }
    h2, h3 { font-weight: 600 !important; color: #1E88E5 !important; margin-top: 2rem !important; }
    .stAlert, .stExpander { border-radius: 10px !important; border-left: 4px solid #1E88E5 !important; box-shadow: 0 2px 8px rgba(0,0,0,0.1) !important; }
    .stButton>button { background: linear-gradient(90deg, #1E88E5 0%, #1976D2 100%) !important; color: white !important; border-radius: 8px !important; padding: 0.75rem 2rem !important; font-weight: 600 !important; border: none !important; box-shadow: 0 4px 12px rgba(30, 136, 229, 0.3) !important; transition: all 0.3s ease !important; }
    .stButton>button:hover { transform: translateY(-2px) !important; box-shadow: 0 6px 16px rgba(30, 136, 229, 0.4) !important; }
    .stDownloadButton>button { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important; border-radius: 8px !important; font-weight: 500 !important; padding: 0.5rem 1.5rem !important; transition: all 0.3s ease !important; }
    .stDownloadButton>button:hover { transform: scale(1.05) !important; box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4) !important; }
    [data-testid="stMetricValue"] { font-size: 2rem !important; font-weight: 700 !important; color: #1E88E5 !important; }
    [data-testid="stFileUploader"] { background: linear-gradient(135deg, #1A1F2E 0%, #2A2F3E 100%) !important; border: 2px dashed #1E88E5 !important; border-radius: 12px !important; padding: 2rem !important; }
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] { background-color: transparent; border-radius: 8px 8px 0 0; padding: 10px 20px; font-weight: 600; }
    .stTabs [aria-selected="true"] { background: linear-gradient(90deg, #1E88E5 0%, #1976D2 100%); color: white !important; }
    .stTextArea textarea, .stTextInput input { border-radius: 8px !important; border: 2px solid #1A1F2E !important; font-family: 'Inter', sans-serif !important; }
    .stTextArea textarea:focus, .stTextInput input:focus { border-color: #1E88E5 !important; box-shadow: 0 0 0 2px rgba(30, 136, 229, 0.2) !important; }
    .stCodeBlock { border-radius: 10px !important; border-left: 4px solid #00BCD4 !important; }
    .stSuccess { background: linear-gradient(135deg, #4CAF50 0%, #66BB6A 100%) !important; border-radius: 8px !important; }
    .stInfo { background: linear-gradient(135deg, #2196F3 0%, #42A5F5 100%) !important; border-radius: 8px !important; }
    section[data-testid="stSidebar"] { background: linear-gradient(180deg, #1A1F2E 0%, #0E1117 100%) !important; }
    .emoji { font-size: 1.2em !important; vertical-align: middle !important; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
    Use this app to run advanced FP&A analysis with a virtual CFO-level AI agent. 
    Provide your business context, metrics, upload financial data, and get comprehensive analysis with visualizations.
""")

api_key = st.secrets.get("GOOGLE_API_KEY", os.getenv("GOOGLE_API_KEY"))
if api_key:
    os.environ["GOOGLE_API_KEY"] = api_key

# Initialize session state
if 'conversation_history' not in st.session_state:
    st.session_state.conversation_history = []
if 'analysis_history' not in st.session_state:
    st.session_state.analysis_history = []
if 'current_response' not in st.session_state:
    st.session_state.current_response = None
if 'uploaded_data' not in st.session_state:
    st.session_state.uploaded_data = None

# Sidebar for history
with st.sidebar:
    st.markdown("### 📜 Analysis History")
    
    if len(st.session_state.analysis_history) > 0:
        for idx, item in enumerate(reversed(st.session_state.analysis_history)):
            with st.expander(f"📊 Analysis #{len(st.session_state.analysis_history) - idx}"):
                st.write(f"**Prompt:** {item['prompt'][:100]}...")
                st.write(f"**Time:** {item['timestamp']}")
                if st.button(f"Load #{len(st.session_state.analysis_history) - idx}", key=f"load_{idx}"):
                    st.session_state.conversation_history = item.get('conversation', [])
                    st.session_state.current_response = item['response']
                    st.rerun()
                if st.button(f"Delete #{len(st.session_state.analysis_history) - idx}", key=f"del_{idx}"):
                    st.session_state.analysis_history.pop(len(st.session_state.analysis_history) - idx - 1)
                    st.rerun()
    else:
        st.info("📄 No history yet. Run an analysis to start!")
    
    if st.button("🗑️ Clear All History"):
        st.session_state.analysis_history = []
        st.session_state.conversation_history = []
        st.session_state.current_response = None
        st.rerun()

if not api_key:
    st.error(
        "GOOGLE_API_KEY is not set. Create a `.env` file in the project folder "
        "and add `GOOGLE_API_KEY=your_key_here`."
    )
else:
    # File upload section
    st.markdown("### 📎 Upload Financial Data (Optional)")
    uploaded_file = st.file_uploader(
        "Upload Excel/CSV file with financial data",
        type=["xlsx", "xls", "csv"],
        help="Upload your financial statements, budgets, or actuals for analysis"
    )
    
    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith('.csv'):
                st.session_state.uploaded_data = pd.read_csv(uploaded_file)
            else:
                st.session_state.uploaded_data = pd.read_excel(uploaded_file)
            st.success(f"✅ File uploaded: {uploaded_file.name} ({len(st.session_state.uploaded_data)} rows)")
            with st.expander("Preview uploaded data"):
                st.dataframe(st.session_state.uploaded_data.head(10))
        except Exception as e:
            st.error(f"Error reading file: {e}")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        audience = st.selectbox(
            "Audience (who will read this output?)",
            ["CEO", "CFO", "Board", "Finance Team", "Operations"],
            index=0,
        )
        st.markdown("### FP&A Prompt")
        prompt = st.text_area(
            "Describe the business, metrics, and what analysis you want.",
            height=200,
            placeholder=(
                "Example:\n"
                "We are a Series B SaaS startup with ARR, churn, CAC, burn, etc.\n"
                "We want a 3-year plan, key KPIs, risks, and recommendations."
            ),
        )
        
        include_charts = st.checkbox("Include data visualizations and metrics", value=True)
        include_formulas = st.checkbox("Generate Google Sheets formulas", value=True)
        
    with col2:
        st.markdown("### 📊 Output & Follow-ups")
        st.markdown("**Follow-up Questions:**")
        
        followup_questions = [
            "🃚 What are the key risks?",
            "💰 What's the financial impact?",
            "✅ What are your recommendations?",
            "🔍 Deep dive into metrics",
            "📊 Show alternative scenarios"
        ]
        
        followup_col1, followup_col2 = st.columns(2)
        selected_followup = None
        
        for i, q in enumerate(followup_questions[:3]):
            if followup_col1.button(q, key=f"followup_{i}"):
                selected_followup = q
        
        for i, q in enumerate(followup_questions[3:], 3):
            if followup_col2.button(q, key=f"followup_{i}"):
                selected_followup = q
    
    run_button = st.button("🚀 Run FP&A Analysis", type="primary", use_container_width=True)
    
    if run_button:
        if not prompt.strip():
            st.warning("Please enter a prompt before running the agent.")
        else:
            with st.spinner("🔄 Running FP&A AI Agent..."):
                try:
                    enhanced_prompt = prompt
                    
                    if st.session_state.uploaded_data is not None:
                        data_summary = f"\\n\\nUploaded Data Summary:\\n"
                        data_summary += f"- Rows: {len(st.session_state.uploaded_data)}\\n"
                        data_summary += f"- Columns: {', '.join(st.session_state.uploaded_data.columns.tolist())}\\n"
                        enhanced_prompt += data_summary
                    
                    if include_charts:
                        enhanced_prompt += "\\n\\nPlease include specific numerical KPIs and metrics in your analysis."
                    
                    if include_formulas:
                        enhanced_prompt += "\\n\\nProvide Google Sheets formulas with IFERROR statements for key calculations."
                    
                    response = run_fp_a_agent(enhanced_prompt, audience=audience)
                    
                    # Add to conversation history
                    st.session_state.conversation_history.append({
                        "role": "user",
                        "content": prompt
                    })
                    st.session_state.conversation_history.append({
                        "role": "assistant",
                        "content": response
                    })
                    
                    st.session_state.current_response = response
                    
                    # Save to analysis history
                    st.session_state.analysis_history.append({
                        "prompt": prompt,
                        "response": response,
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "conversation": st.session_state.conversation_history.copy(),
                        "audience": audience
                    })
                    
                    st.success("✅ Analysis Complete!")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error running FP&A AI Agent: {e}")
    
    # Display current response and follow-ups
    if st.session_state.current_response:
        st.markdown("---")
        st.markdown("### 📄 Analysis Output")
        st.markdown(st.session_state.current_response)
        
        # Download options for OUTPUT
        st.markdown("---")
        st.markdown("### 💾 Download Analysis Results")
        
        # Format numbers in response for better display
        formatted_response = format_numbers_in_text(st.session_state.current_response)
        
        dl_col1, dl_col2, dl_col3, dl_col4 = st.columns(4)
        
        with dl_col1:
            txt_data = formatted_response.encode('utf-8')
            st.download_button(
                label="📄 TXT",
                data=txt_data,
                file_name=f"fpa_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                key="download_txt"
            )
        
        with dl_col2:
            csv_data = convert_response_to_csv(formatted_response)
            st.download_button(
                label="📊 CSV",
                data=csv_data,
                file_name=f"fpa_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                key="download_csv"
            )
        
        with dl_col3:
            try:
                docx_data = convert_response_to_docx(formatted_response)
                st.download_button(
                    label="📃 DOCX",
                    data=docx_data,
                    file_name=f"fpa_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx"
                )
            except:
                st.info("⚠️ Install python-docx for DOCX support")
        
        with dl_col4:
            xlsx_data = convert_response_to_xlsx(formatted_response)
            st.download_button(
                label="📈 XLSX",
                data=xlsx_data,
                file_name=f"fpa_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_xlsx"
            )
        
        # AI Suggestions and Corrections
        st.markdown("---")
        st.markdown("### 🤖 AI Suggestions & Follow-ups")
        
        suggestion_prompt = f"Based on this FP&A analysis, provide 3 smart follow-up questions and 2 corrections/improvements: {st.session_state.current_response[:500]}"
        
        if st.button("🔍 Get Smart Suggestions"):
            with st.spinner("🤔 Generating suggestions..."):
                try:
                    suggestions = run_fp_a_agent(suggestion_prompt, audience=audience)
                    st.info(suggestions)
                    
                    # Add suggestion options
                    st.markdown("**Follow-up Analysis Options:**")
                    sugg_cols = st.columns(2)
                    
                    if sugg_cols[0].button("👍 Accept & Refine Analysis"):
                        st.session_state.conversation_history.append({
                            "role": "assistant",
                            "content": f"Refined: {suggestions}"
                        })
                    if sugg_cols[1].button("🔄 Request Alternative Scenario"):
                        st.write("Request received. Running alternative scenario analysis...")
                except Exception as e:
                    st.error(f"Error generating suggestions: {e}")
        
        # Conversation history display
        if len(st.session_state.conversation_history) > 2:
            st.markdown("---")
            st.markdown("### 🗃️ Conversation History")
            
            for msg in st.session_state.conversation_history:
                if msg["role"] == "user":
                    st.markdown(f"**👤 You:** {msg['content'][:200]}...")
                else:
                    st.markdown(f"**🤖 AI CFO:** {msg['content'][:200]}...")

def format_numbers_in_text(text):
    """Format numbers in text with proper thousand separators and decimals"""
    def format_number(match):
        num = float(match.group())
        if num >= 1000000:
            return f"{num:,.2f}M"
        elif num >= 1000:
            return f"{num:,.2f}K"
        else:
            return f"{num:,.2f}"
    
    pattern = r'\b\d+\.?\d*\b'
    return re.sub(pattern, format_number, text)

def convert_response_to_csv(response):
    """Convert analysis response to CSV format"""
    lines = response.split('\\n')
    csv_data = "Analysis Output\\n\\n"
    for line in lines:
        csv_data += f"{line}\\n"
    return csv_data.encode('utf-8')

def convert_response_to_docx(response):
    """Convert analysis response to DOCX format"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    doc.add_heading('FP&A Analysis Report', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(response)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def convert_response_to_xlsx(response):
    """Convert analysis response to XLSX format"""
    df = pd.DataFrame()
    lines = response.split('\\n')
    
    # Create a structured format
    data = {"FP&A Analysis": lines}
    df = pd.DataFrame(data)
    
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Analysis', index=False)
    buffer.seek(0)
    return buffer.getvalue()
