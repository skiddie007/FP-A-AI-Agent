import os
from dotenv import load_dotenv
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from fp_a_agent import run_fp_a_agent
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt

# Load .env (GOOGLE_API_KEY)
load_dotenv()

st.set_page_config(page_title="FP&A AI Agent", layout="wide")
st.title("FP&A AI Agent – CFO-Grade FP&A with Data Analytics")
st.markdown(
    """
    Use this app to run advanced FP&A analysis with a virtual CFO-level AI agent. 
    Provide your business context, metrics, upload financial data, and get comprehensive analysis with visualizations.
    """
)

api_key = st.secrets.get("GOOGLE_API_KEY", os.getenv("GOOGLE_API_KEY"))
if api_key:
    os.environ["GOOGLE_API_KEY"] = api_key

if not api_key:
    st.error(
        "GOOGLE_API_KEY is not set. Create a `.env` file in the project folder "
        "and add `GOOGLE_API_KEY=your_key_here`."
    )
else:
    # File upload section
    st.markdown("### 📎 Upload Financial Data (Optional)")
    uploaded_file = st.file_uploader(
        "Upload Excel/CSV file with financial data",
        type=["xlsx", "xls", "csv"],
        help="Upload your financial statements, budgets, or actuals for analysis"
    )
    
    uploaded_data = None
    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith('.csv'):
                uploaded_data = pd.read_csv(uploaded_file)
            else:
                uploaded_data = pd.read_excel(uploaded_file)
            st.success(f"✅ File uploaded: {uploaded_file.name} ({len(uploaded_data)} rows)")
            with st.expander("Preview uploaded data"):
                st.dataframe(uploaded_data.head(10))
        except Exception as e:
            st.error(f"Error reading file: {e}")

    col1, col2 = st.columns([1, 1])

    with col1:
        audience = st.selectbox(
            "Audience (who will read this output?)",
            ["CEO", "CFO", "Board", "Finance Team", "Operations"],
            index=0,
        )

        st.markdown("### FP&A Prompt")
        prompt = st.text_area(
            "Describe the business, metrics, and what analysis you want.",
            height=200,
            placeholder=(
                "Example:\n"
                "We are a Series B SaaS startup with ARR, churn, CAC, burn, etc.\n"
                "We want a 3-year plan, key KPIs, risks, and recommendations."
            ),
        )
        
        # Add checkbox for enhanced output
        include_charts = st.checkbox("Include data visualizations and metrics", value=True)
        include_formulas = st.checkbox("Generate Google Sheets formulas", value=True)

        run_button = st.button("Run FP&A Analysis", type="primary")

    with col2:
        st.markdown("### Output")
        output_placeholder = st.empty()

    if run_button:
        if not prompt.strip():
            st.warning("Please enter a prompt before running the agent.")
        else:
            with st.spinner("Running FP&A AI Agent..."):
                try:
                    # Build enhanced prompt
                    enhanced_prompt = prompt
                    
                    if uploaded_data is not None:
                        data_summary = f"\n\nUploaded Data Summary:\n"
                        data_summary += f"- Rows: {len(uploaded_data)}\n"
                        data_summary += f"- Columns: {', '.join(uploaded_data.columns.tolist())}\n"
                        enhanced_prompt += data_summary
                    
                    if include_charts:
                        enhanced_prompt += "\n\nPlease include specific numerical KPIs and metrics in your analysis."
                    
                    if include_formulas:
                        enhanced_prompt += "\n\nProvide Google Sheets formulas with IFERROR statements for key calculations."
                    
                    # Get AI response
                    response = run_fp_a_agent(enhanced_prompt, audience=audience)
                    
                    with col2:
                        output_placeholder.markdown(response)

                                            # Download options
                        st.markdown("---")
                        st.markdown("### 💾 Download Analysis")
                        
                        col_d1, col_d2, col_d3, col_d4 = st.columns(4)
                        
                        # TXT Download
                        with col_d1:
                            txt_data = response
                            st.download_button(
                                label="📄 Download TXT",
                                data=txt_data,
                                file_name=f"FPA_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                                mime="text/plain"
                            )
                        
                        # CSV Download
                        with col_d2:
                            # Create a simple CSV with the response
                            csv_data = f"FP&A Analysis Report\n{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n{response}"
                            st.download_button(
                                label="📊 Download CSV",
                                data=csv_data,
                                file_name=f"FPA_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                mime="text/csv"
                            )
                        
                        # DOCX Download
                        with col_d3:
                            # Create Word document
                            doc = Document()
                            doc.add_heading('FP&A Analysis Report', 0)
                            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                            doc.add_paragraph(response)
                            
                            # Save to bytes
                            docx_buffer = io.BytesIO()
                            doc.save(docx_buffer)
                            docx_buffer.seek(0)
                            
                            st.download_button(
                                label="📝 Download DOCX",
                                data=docx_buffer.getvalue(),
                                file_name=f"FPA_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        # XLSX Download
                        with col_d4:
                            # Create Excel file
                            excel_buffer = io.BytesIO()
                            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                                # Create a dataframe with the analysis
                                df_report = pd.DataFrame({
                                    'FP&A Analysis Report': [response],
                                    'Generated': [datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
                                    'Audience': [audience]
                                })
                                df_report.to_excel(writer, sheet_name='Analysis', index=False)
                            excel_buffer.seek(0)
                            
                            st.download_button(
                                label="📈 Download XLSX",
                                data=excel_buffer.getvalue(),
                                file_name=f"FPA_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                        
                        # Generate sample visualizations if data is uploaded
                        if uploaded_data is not None and include_charts:
                            st.markdown("---")
                            st.markdown("### 📊 Data Visualizations")
                            
                            # Try to detect numeric columns
                            numeric_cols = uploaded_data.select_dtypes(include=['number']).columns.tolist()
                            
                            if len(numeric_cols) > 0:
                                # Sample chart 1: Line chart
                                if len(numeric_cols) >= 1:
                                    fig1 = px.line(uploaded_data, y=numeric_cols[0], 
                                                  title=f"{numeric_cols[0]} Trend")
                                    st.plotly_chart(fig1, use_container_width=True)
                                
                                # Sample chart 2: Bar chart
                                if len(numeric_cols) >= 2:
                                    fig2 = px.bar(uploaded_data.head(10), y=numeric_cols[:min(2, len(numeric_cols))],
                                                 title="Key Metrics Comparison")
                                    st.plotly_chart(fig2, use_container_width=True)
                        
                        # Display Google Sheets formulas if requested
                        if include_formulas:
                            st.markdown("---")
                            st.markdown("### 📐 Google Sheets Formulas")
                            st.code("""
# Revenue Growth Rate
=IFERROR((B2-B1)/B1, 0)

# EBITDA Margin
=IFERROR(EBITDA/Revenue, 0)

# Burn Rate (Monthly)
=IFERROR(SUM(Expenses)/COUNT(Months), 0)

# Customer Acquisition Cost (CAC)
=IFERROR(Sales_Marketing_Spend/New_Customers, 0)

# Lifetime Value (LTV)
=IFERROR(ARPU*Customer_Lifetime_Months, 0)

# LTV/CAC Ratio
=IFERROR(LTV/CAC, 0)

# Cash Runway (Months)
=IFERROR(Cash_Balance/Monthly_Burn, 0)

# Gross Margin
=IFERROR((Revenue-COGS)/Revenue, 0)

# Operating Cash Flow
=IFERROR(EBITDA-CapEx-Working_Capital_Change, 0)

# Return on Equity (ROE)
=IFERROR(Net_Income/Shareholders_Equity, 0)
                            """, language="excel")
                            
                except Exception as e:
                    st.error(f"Error running FP&A AI Agent: {e}")
